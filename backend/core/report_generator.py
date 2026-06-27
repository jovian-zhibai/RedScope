"""Report generation engine: generates Word/PDF pentest reports."""

from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import select
from sqlalchemy.orm import Session
from backend.models.project import Project
from backend.models.finding import Finding
from backend.models.asset import Asset
from backend.models.operational import AttackTimeline, CompromisedHost
from backend.config import get_settings


SEVERITY_CN = {"critical": "严重", "high": "高危", "medium": "中危", "low": "低危", "info": "信息"}
SEVERITY_COLOR = {"critical": RGBColor(0xF8, 0x51, 0x49), "high": RGBColor(0xD2, 0x99, 0x22),
                  "medium": RGBColor(0x58, 0xA6, 0xFF), "low": RGBColor(0x3F, 0xB9, 0x50),
                  "info": RGBColor(0x8B, 0x94, 0x9E)}


class ReportGenerator:
    def __init__(self, db: Session, project_id: int):
        self.db = db
        self.project = db.get(Project, project_id)
        self.findings = list(db.execute(
            select(Finding).where(Finding.project_id == project_id, Finding.is_false_positive == False, Finding.deleted_at == None)
            .order_by(Finding.severity)
        ).scalars().all())
        self.assets = list(db.execute(select(Asset).where(Asset.project_id == project_id, Asset.deleted_at == None)).scalars().all())
        self.timeline = list(db.execute(
            select(AttackTimeline).where(AttackTimeline.project_id == project_id)
            .order_by(AttackTimeline.timestamp)
        ).scalars().all())
        self.settings = get_settings()

    def generate_docx(self, output_path: str | None = None) -> str:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(10)

        self._add_cover(doc)
        self._add_summary(doc)
        self._add_scope(doc)
        self._add_findings_section(doc)
        self._add_statistics(doc)
        self._add_timeline_section(doc)
        self._add_remediation(doc)

        if not output_path:
            output_dir = Path(self.settings.scan_output_dir) / "reports"
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = str(output_dir / f"report_{self.project.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

        doc.save(output_path)
        return output_path

    def _add_cover(self, doc: Document):
        doc.add_paragraph("")
        doc.add_paragraph("")
        title = doc.add_heading(level=0)
        run = title.add_run("渗透测试报告")
        run.font.size = Pt(28)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(f"\n{self.project.name}\n").font.size = Pt(16)
        if self.project.client_name:
            subtitle.add_run(f"\n客户: {self.project.client_name}\n").font.size = Pt(12)
        subtitle.add_run(f"\n报告日期: {datetime.now().strftime('%Y年%m月%d日')}\n").font.size = Pt(12)
        subtitle.add_run("\n\n机密文件 - 仅限授权人员查阅").font.color.rgb = RGBColor(0xF8, 0x51, 0x49)
        doc.add_page_break()

    def _add_summary(self, doc: Document):
        doc.add_heading("一、测试概述", level=1)
        sev_counts = {}
        for f in self.findings:
            sev_counts[f.severity] = sev_counts.get(f.severity, 0) + 1

        doc.add_paragraph(
            f"本次渗透测试共发现 {len(self.findings)} 个安全漏洞，"
            f"其中严重 {sev_counts.get('critical', 0)} 个，"
            f"高危 {sev_counts.get('high', 0)} 个，"
            f"中危 {sev_counts.get('medium', 0)} 个，"
            f"低危 {sev_counts.get('low', 0)} 个，"
            f"信息 {sev_counts.get('info', 0)} 个。"
        )
        doc.add_paragraph(f"测试资产总数: {len(self.assets)} 个")
        if self.project.auth_start_date and self.project.auth_end_date:
            doc.add_paragraph(f"测试周期: {self.project.auth_start_date} 至 {self.project.auth_end_date}")

    def _add_scope(self, doc: Document):
        doc.add_heading("二、测试范围", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = "主机/域名"
        headers[1].text = "端口"
        headers[2].text = "应用"
        headers[3].text = "重要性"
        for asset in self.assets[:50]:
            row = table.add_row().cells
            row[0].text = asset.host or ""
            row[1].text = str(asset.port or "")
            row[2].text = asset.application or asset.server or ""
            row[3].text = asset.importance or ""

    def _add_findings_section(self, doc: Document):
        doc.add_heading("三、漏洞详情", level=1)
        for i, f in enumerate(self.findings, 1):
            doc.add_heading(f"3.{i} {f.title}", level=2)

            p = doc.add_paragraph()
            run = p.add_run(f"风险等级: {SEVERITY_CN.get(f.severity, f.severity)}")
            run.font.color.rgb = SEVERITY_COLOR.get(f.severity, RGBColor(0, 0, 0))
            run.bold = True
            if f.cvss_score:
                p.add_run(f"  |  CVSS: {f.cvss_score}")

            if f.description:
                doc.add_heading("漏洞描述", level=3)
                doc.add_paragraph(f.description)
            if f.detail:
                doc.add_heading("复现步骤", level=3)
                doc.add_paragraph(f.detail)
            if f.solution:
                doc.add_heading("修复建议", level=3)
                doc.add_paragraph(f.solution)
            if f.evidence and isinstance(f.evidence, dict):
                if f.evidence.get("request"):
                    doc.add_heading("请求证据", level=3)
                    doc.add_paragraph(str(f.evidence["request"])[:2000], style='No Spacing')

    def _add_statistics(self, doc: Document):
        doc.add_heading("四、漏洞统计", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "风险等级"
        table.rows[0].cells[1].text = "数量"
        for sev in ["critical", "high", "medium", "low", "info"]:
            count = sum(1 for f in self.findings if f.severity == sev)
            if count > 0:
                row = table.add_row().cells
                row[0].text = SEVERITY_CN.get(sev, sev)
                row[1].text = str(count)

    def _add_timeline_section(self, doc: Document):
        if not self.timeline:
            return
        doc.add_heading("五、攻击时间线", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        h = table.rows[0].cells
        h[0].text = "时间"
        h[1].text = "阶段"
        h[2].text = "操作"
        h[3].text = "目标"
        for t in self.timeline:
            row = table.add_row().cells
            row[0].text = t.timestamp.strftime("%m-%d %H:%M") if t.timestamp else ""
            row[1].text = t.phase or ""
            row[2].text = t.action or ""
            row[3].text = t.target_host or ""

    def _add_remediation(self, doc: Document):
        doc.add_heading("六、修复优先级建议", level=1)
        critical_high = [f for f in self.findings if f.severity in ("critical", "high")]
        if critical_high:
            doc.add_paragraph("以下漏洞建议在 7 天内完成修复:")
            for f in critical_high:
                doc.add_paragraph(f"• [{SEVERITY_CN.get(f.severity)}] {f.title}", style='List Bullet')
        medium = [f for f in self.findings if f.severity == "medium"]
        if medium:
            doc.add_paragraph("\n以下漏洞建议在 30 天内完成修复:")
            for f in medium:
                doc.add_paragraph(f"• [中危] {f.title}", style='List Bullet')
